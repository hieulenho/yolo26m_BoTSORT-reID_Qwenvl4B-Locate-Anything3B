from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

FORBIDDEN_PATTERNS = (
    r"lorem ipsum",
    r"\bTBD\b",
    r"\bTODO\b",
    r"placeholder",
)

EXPECTED_HEADINGS = [
    "I. Introduction",
    "II. Related Work",
    "III. Method",
    "IV. Experimental Protocol",
    "V. Results",
    "VI. Error Analysis and Ablation",
    "VII. Discussion",
    "VIII. Limitations",
    "IX. Conclusion",
    "References",
]
EXPECTED_TABLE_LABELS = [
    "I",
    "II",
    "III",
    "IV",
    "V",
    "VI",
    "VII",
    "VIII",
    "IX",
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def pdf_text(path: Path) -> tuple[int, str]:
    reader = PdfReader(str(path))
    return len(reader.pages), "\n".join(page.extract_text() or "" for page in reader.pages)


def _resolve_pdf_object(value: Any) -> Any:
    return value.get_object() if hasattr(value, "get_object") else value


def _descriptor_embeds_font(descriptor: Any) -> bool:
    descriptor = _resolve_pdf_object(descriptor)
    return bool(
        descriptor
        and any(key in descriptor for key in ("/FontFile", "/FontFile2", "/FontFile3"))
    )


def _font_is_embedded(font: Any) -> bool:
    font = _resolve_pdf_object(font)
    subtype = str(font.get("/Subtype", ""))
    if subtype == "/Type3":
        return True
    descendants = _resolve_pdf_object(font.get("/DescendantFonts", [])) or []
    if descendants:
        return all(
            _descriptor_embeds_font(_resolve_pdf_object(item).get("/FontDescriptor"))
            for item in descendants
        )
    return _descriptor_embeds_font(font.get("/FontDescriptor"))


def inspect_pdf_fonts(path: Path) -> list[dict[str, object]]:
    reader = PdfReader(str(path))
    records: dict[tuple[str, str], dict[str, object]] = {}
    for page in reader.pages:
        resources = _resolve_pdf_object(page.get("/Resources", {})) or {}
        fonts = _resolve_pdf_object(resources.get("/Font", {})) or {}
        for font_ref in fonts.values():
            font = _resolve_pdf_object(font_ref)
            key = (str(font.get("/BaseFont", "unknown")), str(font.get("/Subtype", "unknown")))
            records[key] = {
                "base_font": key[0],
                "subtype": key[1],
                "embedded": _font_is_embedded(font),
            }
    return [records[key] for key in sorted(records)]


def xml_text(payload: bytes) -> str:
    root = ElementTree.fromstring(payload)
    return " ".join(node.text or "" for node in root.iter() if node.tag.endswith("}t"))


def inspect_pptx(path: Path) -> dict[str, object]:
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        slide_names = sorted(
            name for name in names if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
        )
        notes_names = sorted(
            name for name in names if re.fullmatch(r"ppt/notesSlides/notesSlide\d+\.xml", name)
        )
        slide_text = "\n".join(xml_text(archive.read(name)) for name in slide_names)
        notes_texts = [xml_text(archive.read(name)) for name in notes_names]

    notes_with_sources = sum(
        "[Sources]" in notes and "[/Sources]" in notes for notes in notes_texts
    )
    return {
        "slide_count": len(slide_names),
        "notes_count": len(notes_names),
        "notes_with_sources": notes_with_sources,
        "text": slide_text,
    }


def _close(value: float, expected: float, tolerance: float = 0.002) -> bool:
    return abs(value - expected) <= tolerance


def _section_columns(section: Any) -> tuple[int, int]:
    columns = section._sectPr.find(qn("w:cols"))
    if columns is None:
        return 1, 720
    count = int(columns.get(qn("w:num"), "1"))
    gap = int(columns.get(qn("w:space"), "720"))
    return count, gap


def _paragraph_font_size(paragraph: Any) -> float | None:
    sizes = [run.font.size.pt for run in paragraph.runs if run.font.size is not None]
    return sizes[0] if sizes else None


def _word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text, flags=re.UNICODE))


def inspect_docx_format(path: Path) -> dict[str, object]:
    document = Document(path)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    paragraph_by_prefix = {paragraph.text.strip(): paragraph for paragraph in paragraphs}
    abstract = next(
        paragraph for paragraph in paragraphs if paragraph.text.strip().startswith("Abstract")
    )
    keywords = next(
        paragraph for paragraph in paragraphs if paragraph.text.strip().startswith("Keywords")
    )
    title = paragraphs[0]
    author = paragraphs[1]
    affiliation = paragraphs[2]

    abstract_body = re.sub(r"^Abstract\s*[—-]\s*", "", abstract.text.strip())
    keyword_body = re.sub(r"^Keywords\s*[—-]\s*", "", keywords.text.strip()).rstrip(".")
    keyword_items = [item.strip() for item in keyword_body.split(",") if item.strip()]
    heading_order = [text for text in EXPECTED_HEADINGS if text in paragraph_by_prefix]
    figure_numbers = [
        int(match.group(1))
        for paragraph in paragraphs
        if (match := re.match(r"^Fig\.\s*(\d+)\.", paragraph.text.strip()))
    ]
    table_labels = [
        match.group(1)
        for paragraph in paragraphs
        if (match := re.match(r"^TABLE\s+([IVX]+)\b", paragraph.text.strip()))
    ]
    reference_numbers = [
        int(match.group(1))
        for paragraph in paragraphs
        if (match := re.match(r"^\[(\d+)\]", paragraph.text.strip()))
    ]

    section_records = []
    for section in document.sections:
        count, gap = _section_columns(section)
        section_records.append(
            {
                "page_width_inches": round(section.page_width.inches, 4),
                "page_height_inches": round(section.page_height.inches, 4),
                "top_margin_inches": round(section.top_margin.inches, 4),
                "bottom_margin_inches": round(section.bottom_margin.inches, 4),
                "left_margin_inches": round(section.left_margin.inches, 4),
                "right_margin_inches": round(section.right_margin.inches, 4),
                "columns": count,
                "column_gap_twips": gap,
            }
        )

    header_footer_text = " ".join(
        paragraph.text.strip()
        for section in document.sections
        for container in (section.header, section.footer)
        for paragraph in container.paragraphs
        if paragraph.text.strip()
    )
    with zipfile.ZipFile(path) as archive:
        field_instructions = []
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            root = ElementTree.fromstring(archive.read(name))
            field_instructions.extend(
                (node.text or "").strip()
                for node in root.iter()
                if node.tag.endswith("}instrText")
            )
    has_page_field = any(
        re.search(r"\b(?:PAGE|NUMPAGES)\b", instruction, flags=re.IGNORECASE)
        for instruction in field_instructions
    )

    all_letter = all(
        _close(record["page_width_inches"], 8.5)
        and _close(record["page_height_inches"], 11.0)
        for record in section_records
    )
    margins_match = all(
        _close(record["top_margin_inches"], 0.75)
        and _close(record["bottom_margin_inches"], 1.0)
        and _close(record["left_margin_inches"], 0.625)
        and _close(record["right_margin_inches"], 0.625)
        for record in section_records
    )
    body_columns_match = len(section_records) >= 2 and section_records[1]["columns"] == 2
    body_gap_matches = len(section_records) >= 2 and section_records[1]["column_gap_twips"] == 360

    checks = {
        "letter_page_size": all_letter,
        "ieee_conference_margins": margins_match,
        "two_column_body": body_columns_match,
        "quarter_inch_column_gap": body_gap_matches,
        "title_is_24_pt": _paragraph_font_size(title) == 24.0,
        "author_is_11_pt": _paragraph_font_size(author) == 11.0,
        "affiliation_is_10_pt": _paragraph_font_size(affiliation) == 10.0,
        "abstract_is_one_paragraph_and_at_most_250_words": 1 <= _word_count(abstract_body) <= 250,
        "abstract_is_9_pt": _paragraph_font_size(abstract) == 9.0,
        "keyword_count_is_3_to_5": 3 <= len(keyword_items) <= 5,
        "section_order_is_complete": heading_order == EXPECTED_HEADINGS,
        "figure_captions_are_sequential": figure_numbers == list(range(1, 7)),
        "table_captions_are_sequential": table_labels == EXPECTED_TABLE_LABELS,
        "references_are_sequential": reference_numbers == list(range(1, 21)),
        "no_header_or_footer_text": not header_footer_text,
        "no_page_number_field": not has_page_field,
    }
    return {
        "checks": checks,
        "sections": section_records,
        "abstract_word_count": _word_count(abstract_body),
        "keywords": keyword_items,
        "headings": heading_order,
        "figure_numbers": figure_numbers,
        "table_labels": table_labels,
        "reference_numbers": reference_numbers,
        "field_instructions": field_instructions,
    }


def missing_terms(text: str, terms: list[str]) -> list[str]:
    normalized = " ".join(text.split())
    return [term for term in terms if term not in normalized]


def forbidden_matches(text: str) -> list[str]:
    return [
        pattern for pattern in FORBIDDEN_PATTERNS if re.search(pattern, text, flags=re.IGNORECASE)
    ]


def artifact_record(path: Path) -> dict[str, object]:
    return {
        "path": str(path),
        "bytes": path.stat().st_size,
        "sha256": sha256(path),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Verify the final research paper and deck.")
    parser.add_argument("--repo-root", type=Path, default=Path(__file__).resolve().parents[2])
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("outputs/reports/research_artifacts.verification.json"),
    )
    args = parser.parse_args()
    root = args.repo_root.resolve()
    output = args.output if args.output.is_absolute() else root / args.output

    paper_docx = root / "outputs/reports/research_paper/locate_then_reason_ieee_paper.docx"
    paper_pdf = root / "outputs/reports/research_paper/locate_then_reason_ieee_paper.pdf"
    deck_pptx = root / "outputs/reports/research_presentation/locate_then_reason_defense.pptx"
    deck_pdf = root / "outputs/reports/research_presentation/locate_then_reason_defense.pdf"
    paths = [paper_docx, paper_pdf, deck_pptx, deck_pdf]

    missing_files = [str(path) for path in paths if not path.exists()]
    if missing_files:
        raise FileNotFoundError(f"Missing final artifacts: {missing_files}")

    paper_pages, paper_body = pdf_text(paper_pdf)
    deck_pdf_pages, deck_pdf_body = pdf_text(deck_pdf)
    deck = inspect_pptx(deck_pptx)
    ieee_format = inspect_docx_format(paper_docx)
    paper_fonts = inspect_pdf_fonts(paper_pdf)

    checks = {
        "paper_page_count_within_project_target": 4 <= paper_pages <= 8,
        "paper_ieee_format_checks_pass": all(ieee_format["checks"].values()),
        "paper_pdf_fonts_are_embedded": bool(paper_fonts)
        and all(record["embedded"] for record in paper_fonts),
        "deck_slide_count_is_15": deck["slide_count"] == 15,
        "deck_pdf_page_count_is_15": deck_pdf_pages == 15,
        "every_slide_has_speaker_notes": deck["notes_count"] == deck["slide_count"],
        "every_slide_note_has_sources": deck["notes_with_sources"] == deck["slide_count"],
        "paper_required_terms_present": not missing_terms(
            paper_body,
            ["TrackTrack", "LocateAnything", "Qwen3-VL", "87.605", "96.399", "49.52"],
        ),
        "deck_required_terms_present": not missing_terms(
            str(deck["text"]),
            ["TrackTrack", "LocateAnything", "Qwen3-VL-4B", "71,06", "87,61", "87,5%"],
        ),
        "paper_has_no_placeholders": not forbidden_matches(paper_body),
        "deck_has_no_placeholders": not forbidden_matches(str(deck["text"])),
        "deck_pdf_has_no_placeholders": not forbidden_matches(deck_pdf_body),
    }
    status = "ok" if all(checks.values()) else "failed"

    report = {
        "status": status,
        "checks": checks,
        "paper": {
            **artifact_record(paper_docx),
            "pdf": artifact_record(paper_pdf),
            "page_count": paper_pages,
            "ieee_format": ieee_format,
            "pdf_fonts": paper_fonts,
            "missing_required_terms": missing_terms(
                paper_body,
                ["TrackTrack", "LocateAnything", "Qwen3-VL", "87.605", "96.399", "49.52"],
            ),
            "forbidden_matches": forbidden_matches(paper_body),
        },
        "presentation": {
            **artifact_record(deck_pptx),
            "pdf": artifact_record(deck_pdf),
            "slide_count": deck["slide_count"],
            "notes_count": deck["notes_count"],
            "notes_with_sources": deck["notes_with_sources"],
            "missing_required_terms": missing_terms(
                str(deck["text"]),
                ["TrackTrack", "LocateAnything", "Qwen3-VL-4B", "71,06", "87,61", "87,5%"],
            ),
            "forbidden_matches": forbidden_matches(str(deck["text"])),
        },
        "source_scope": {
            "official_gt_metrics": [
                "SportsMOT detector validation",
                "SportsMOT TrackEval over 30 sequences",
                "UA-DETRAC MVI_40774 ignore-aware tracking",
            ],
            "limited_gt_metrics": [
                "LocateAnything-to-Qwen semantic evaluation: 8 scored tracks out of 24 processed",
            ],
            "heuristic_only": [
                "IDSW cause taxonomy",
            ],
        },
    }

    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({"status": status, "output": str(output)}, indent=2))
    if status != "ok":
        raise SystemExit(1)


if __name__ == "__main__":
    main()
