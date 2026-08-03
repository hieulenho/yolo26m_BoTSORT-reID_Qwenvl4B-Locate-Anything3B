from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TITLE = (
    "Locate-Then-Reason: Identity-Stable Semantic Multi-Object Tracking "
    "on Resource-Constrained Hardware"
)


def _set_cell_margins(
    cell, top: int = 55, start: int = 55, bottom: int = 55, end: int = 55
) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_columns(section, count: int, space_twips: int = 360) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))


def _configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def _style_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_section_heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.small_caps = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_subheading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 0.95
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_equation(document: Document, text: str, number: int) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.cell(0, 0).width = Inches(2.65)
    table.cell(0, 1).width = Inches(0.38)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left.add_run(text).italic = True
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(f"({number})")
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(12)


def add_table(
    document: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
) -> None:
    label, description = caption.split(". ", 1)
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(12)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.line_spacing = Pt(10)
    cap.paragraph_format.keep_with_next = True
    label_run = cap.add_run(label.upper())
    label_run.font.name = "Times New Roman"
    label_run.font.size = Pt(8)
    label_run.font.small_caps = True
    label_run.add_break()
    description_run = cap.add_run(description.upper())
    description_run.font.name = "Times New Roman"
    description_run.font.size = Pt(8)
    description_run.font.small_caps = True

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        header_run = paragraph.add_run(header)
        header_run.bold = True
        header_run.font.name = "Times New Roman"
        header_run.font.size = Pt(8)
        if widths:
            cell.width = Inches(widths[index])

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            value_run = paragraph.add_run(str(value))
            value_run.font.name = "Times New Roman"
            value_run.font.size = Pt(8)
            if widths:
                cell.width = Inches(widths[index])

    # Keep compact IEEE tables together when they fit in one column. This
    # avoids leaving a caption and one data row at the foot of a page.
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(
    document: Document, image_path: Path, caption: str, width_inches: float = 3.15
) -> None:
    if not image_path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.line_spacing = Pt(12)
    run = cap.add_run(caption)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)


def add_reference(document: Document, index: int, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.line_spacing = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"[{index}] {text}")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)


def build_document(repo_root: Path, output_docx: Path) -> None:
    traffic = json.loads(
        (repo_root / "docs/benchmarks/traffic_quality/traffic_quality_summary.json").read_text(
            encoding="utf-8"
        )
    )
    research = json.loads(
        (repo_root / "docs/benchmarks/research_final/research_summary.json").read_text(
            encoding="utf-8"
        )
    )

    document = Document()
    _style_document(document)
    first_section = document.sections[0]
    _configure_section(first_section)
    _set_columns(first_section, 1)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.0
    title_lines = (
        "Locate-Then-Reason: Identity-Stable Semantic",
        "Multi-Object Tracking on Resource-Constrained Hardware",
    )
    for index, line in enumerate(title_lines):
        title_run = title.add_run(line)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(24)
        if index == 0:
            title_run.add_break()

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.space_before = Pt(18)
    author.paragraph_format.space_after = Pt(2)
    author.paragraph_format.line_spacing = Pt(12)
    author_run = author.add_run("Hieu Le Nho")
    author_run.font.name = "Times New Roman"
    author_run.font.size = Pt(11)

    affiliation = document.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliation.paragraph_format.space_after = Pt(0)
    affiliation.paragraph_format.line_spacing = Pt(12)
    affiliation_run = affiliation.add_run(
        "Artificial Intelligence Department, Institute of Military Science and Technology, Vietnam"
    )
    affiliation_run.font.name = "Times New Roman"
    affiliation_run.font.size = Pt(10)

    body_section = document.add_section(WD_SECTION.CONTINUOUS)
    _configure_section(body_section)
    _set_columns(body_section, 2, 360)

    abstract = document.add_paragraph()
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.first_line_indent = Inches(0.19)
    abstract.paragraph_format.space_after = Pt(10)
    abstract.paragraph_format.line_spacing = Pt(12)
    prefix = abstract.add_run("Abstract—")
    prefix.bold = True
    prefix.italic = True
    body = abstract.add_run(
        "Online multi-object tracking produces stable spatial identities, whereas large vision-language models provide open-ended semantic descriptions but are too expensive to execute on every frame. This paper presents a resource-aware locate-then-reason pipeline that separates an immediate foreground path from asynchronous semantic refinement. YOLO26 supplies detections; TrackTrack with a convolutional-neural-network appearance encoder performs identity association; LocateAnything-3B localizes an object inside accumulated track evidence; and eight-bit Qwen3-VL-4B assigns hierarchical labels that are fused over time with explicit unknown rejection. The evaluation spans football, traffic, wildlife, microscopy, repeated runtime tests, and a physical webcam. On the thirty-sequence SportsMOT validation protocol, TrackTrack reaches 71.058 Higher Order Tracking Accuracy and 91.511 Multiple Object Tracking Accuracy. On an ignore-aware 750-frame traffic sequence, YOLO26s with TrackTrack appearance matching obtains 87.605 Higher Order Tracking Accuracy, a 96.399 identification F1 score, no identity switches, and 14.16 processed frames per second; a YOLO26s-ByteTrack speed profile reaches 49.52 frames per second. LocateAnything followed by Qwen attains 87.5% parent-class accuracy and 85.5% macro-averaged F1 on the eight reviewed semantic ground-truth tracks available in the sequence. Rich labels arrive asynchronously at 25.51 seconds per processed track on an eight-gigabyte mobile graphics processor, so immediate video remains detector-and-tracker driven while semantic labels are cached and applied to later frames."
    )
    for run in (prefix, body):
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.bold = True

    keywords = document.add_paragraph()
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.first_line_indent = Inches(0.19)
    keywords.paragraph_format.space_after = Pt(6)
    keywords.paragraph_format.line_spacing = Pt(12)
    kw_prefix = keywords.add_run("Keywords—")
    kw_prefix.bold = True
    kw_prefix.italic = True
    kw_body = keywords.add_run(
        "multi-object tracking, open-vocabulary perception, real-time video analytics, vision-language models, visual grounding."
    )
    for run in (kw_prefix, kw_body):
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.bold = True
        run.italic = True

    add_section_heading(document, "I. Introduction")
    add_body(
        document,
        "Tracking-by-detection systems can draw boxes and preserve object identities at video rate, but their class vocabulary is limited by the detector. A large vision-language model (VLM) can describe a tracked object more deeply—for example, distinguishing a sedan from a generic car or a teacher from a generic person—but per-frame multimodal inference is incompatible with low-latency deployment on an 8-GB mobile graphics processing unit (GPU). The central design question is therefore not whether to replace tracking with a VLM, but how to let a VLM enrich a stable identity without blocking the video loop.",
    )
    add_body(
        document,
        "We address this question with a split architecture. The foreground path performs capture, detection, association, rendering, and bounded queue bookkeeping. A semantic worker consumes selected multi-time evidence only when a track is sufficiently visible, persistent, or uncertain. LocateAnything performs query-conditioned spatial verification before Qwen receives the panel and refined crop. Accepted labels are fused in a cache keyed by track identity; weak or contradictory evidence yields unknown rather than a forced class. Fig. 1 summarizes the complete data flow.",
    )
    add_body(
        document,
        "The contribution is a reproducible systems study rather than a new detector or foundation model. First, it defines a modular semantic multi-object tracking (MOT) architecture whose foreground path remains operational when the VLM is disabled or delayed. Second, it integrates track-aware association, convolutional neural network (CNN) appearance cues, re-identification (ReID), localization-first evidence, temporal semantic fusion, and rejection. Third, it evaluates the design under common detections, official ground truth, ignore regions, repeated runtime profiles, and physical webcam capture. Finally, it distinguishes official MOT metrics from a heuristic five-way diagnostic taxonomy for identity failures, preventing causal labels from being mistaken for benchmark truth.",
    )

    add_figure(
        document,
        repo_root / "docs/benchmarks/research_final/figures/research_pipeline_architecture.png",
        "Fig. 1. System overview. The foreground tracker is decoupled from the asynchronous locate-then-reason semantic path.",
    )

    add_section_heading(document, "II. Related Work")
    add_subheading(document, "A. Detection and open-vocabulary perception")
    add_body(
        document,
        "YOLO26 provides end-to-end real-time detection heads and multiple model scales [1]. For known football classes, supervised fine-tuning is retained because its task-specific recall is substantially higher than generic pretrained weights. YOLOE extends the YOLO family toward promptable open-vocabulary detection [2]; in this system it is an optional detector for classes outside a closed model, not a substitute for evaluation against domain ground truth.",
    )
    add_subheading(document, "B. Online multi-object tracking")
    add_body(
        document,
        "SORT [5], DeepSORT [6], ByteTrack [7], BoT-SORT [8], OC-SORT [9], and Deep OC-SORT [10] expose different speed, motion, appearance, and camera-motion trade-offs. TrackTrack [4] reframes association from the track perspective and limits redundant initialization through track-aware initialization. Our implementation adds optional CNN-ReID embeddings and temporal class stabilization, while retaining TrackEval [12] as the metric implementation and HOTA [11] as the primary balanced measure.",
    )
    add_subheading(document, "C. Grounding and semantic reasoning")
    add_body(
        document,
        "Qwen3-VL supports multi-image and video reasoning [14], but autoregressive generation is expensive for dense localization. LocateAnything introduces parallel box decoding for language-conditioned localization [15]. Their roles are complementary: LocateAnything answers where the query-compatible evidence is inside a crop, and Qwen answers what the localized entity most plausibly is. The pipeline therefore preserves box generation and identity association in dedicated real-time models while reserving VLM compute for semantic refinement.",
    )

    add_section_heading(document, "III. Method")
    add_subheading(document, "A. Foreground path")
    add_body(
        document,
        "Each captured frame receives a monotonic timestamp and enters a bounded ring buffer. The detector emits a box, confidence, and base class. The tracker predicts each active state, computes candidate associations, updates matched tracks, retains unmatched tracks for a bounded age, and initializes only plausible unmatched detections. The renderer immediately displays the best known label; before semantic acceptance this is the detector class or a pending state.",
    )
    add_equation(
        document,
        "Cᵢⱼ = λᵢdIoU(i,j) + λₐdapp(i,j) + λₛ(1 − sⱼ)",
        1,
    )
    add_body(
        document,
        "Here dIoU = 1 − IoU, dapp = 1 − cos(eᵢ,eⱼ), e denotes the CNN appearance embedding, and s is the detection confidence. A direction-consistency penalty is applied as a gated auxiliary term. Appearance is gated by spatial plausibility rather than used as an unconditional nearest-neighbor rule. Detector class is stabilized over the track history and is not used as a hard association gate in the selected traffic profile; the controlled ablation showed that a hard gate increased fragmentation and identity switches.",
    )
    add_subheading(document, "B. Track evidence manager")
    add_body(
        document,
        "The evidence manager stores observations by track identity and ranks candidate crops by box area, sharpness, visibility, detector confidence, boundary truncation, and temporal diversity. A semantic job contains an evidence panel, one or more representative crops, the base class, track age, motion diagnostics, and an optional task vocabulary. Queue admission is bounded by pending-event capacity and a cooldown so dense scenes cannot starve the foreground path.",
    )
    add_subheading(document, "C. Locate-then-reason semantic path")
    add_body(
        document,
        "For each selected track, LocateAnything receives the evidence image and a natural-language grounding query derived from the task vocabulary. The returned region is validated, clipped to image bounds, and rejected when its geometry is implausible. Qwen then receives two complementary views: the context panel and the LocateAnything-refined crop. A constrained JSON schema requests coarse class, fine label, visible attributes, confidence, and an abstention reason. Both foundation models are loaded sequentially in 8-bit mode so their peak memory is not additive.",
    )
    add_subheading(document, "D. Temporal fusion and unknown rejection")
    add_body(
        document,
        "Semantic evidence is accumulated per track rather than overwritten by the latest answer. For class c, the cache score follows an exponential update,",
    )
    add_equation(document, "Sₜ(c) = ρSₜ₋₁(c) + (1 − ρ)qₜ(c)", 2)
    add_body(
        document,
        "where qₜ is the normalized confidence from the current event. A label is accepted only if its score exceeds a class threshold, its margin over the runner-up exceeds a second threshold, and repeated observations satisfy the minimum support count. Otherwise, the rendered semantic state remains pending or unknown. This prevents one blurred crop from permanently relabeling an identity.",
    )

    add_section_heading(document, "IV. Experimental Protocol")
    add_subheading(document, "A. Datasets and ground truth")
    add_body(
        document,
        "SportsMOT [3] supplies the football detector split and a 30-sequence, 20,171-frame tracking validation protocol. UA-DETRAC [16] supplies vehicle boxes and identity ground truth; four ignore regions are applied symmetrically to predictions and annotations. AnimalTrack [17] tests open-vocabulary wildlife tracking, and the Cell Tracking Challenge [18] tests a microscopy adapter with and without CLAHE. Semantic accuracy is reported only where reviewed official parent-class labels exist.",
    )
    add_subheading(document, "B. Metrics")
    add_body(
        document,
        "Detection uses precision, recall, mean average precision at 0.5 intersection over union (mAP50), and mean average precision from 0.5 to 0.95 (mAP50–95). Tracking uses Higher Order Tracking Accuracy (HOTA), detection accuracy (DetA), association accuracy (AssA), Multiple Object Tracking Accuracy (MOTA), identification F1 score (IDF1), identity switches (IDSW), false positives (FP), false negatives (FN), and fragmentation. Runtime reports processed frames per second, end-to-end frames per second, 95th-percentile latency, frame drop, and peak memory. Semantic evaluation reports accuracy, macro-averaged F1, coverage, hallucination rate, fine-label accuracy, unknown-rejection F1, and seconds per processed track. A metric is marked not applicable when the ground truth does not support it.",
    )
    add_subheading(document, "C. Hardware and reproducibility")
    hardware = traffic["hardware"]
    add_body(
        document,
        f"All local measurements were recorded on an {hardware['gpu_name']} with 8.0 GiB video memory (VRAM), {hardware['logical_cpu_count']} logical central-processing-unit threads, 15.7 GiB system memory, PyTorch {hardware['torch']}, CUDA {hardware['cuda_runtime']}, and Ultralytics {hardware['ultralytics']}. The report builder records source artifact paths and hashes. Runtime profiles use three foreground repetitions where available; VLM timing is measured separately from the video loop. Table I summarizes the detector operating points used in the study.",
    )

    detector_rows = []
    for item in research["detectors"]:
        detector_rows.append(
            [
                item["display_name"],
                f"{item['precision']:.3f}",
                f"{item['recall']:.3f}",
                f"{item['map50']:.3f}",
                f"{item['map50_95']:.3f}",
                f"{item['detector_fps']:.1f}",
            ]
        )
    add_table(
        document,
        "Table I. SportsMOT detector comparison at 640 pixels",
        ["Detector", "P", "R", "AP50", "AP", "FPS"],
        detector_rows,
        [1.18, 0.30, 0.30, 0.36, 0.36, 0.34],
    )

    add_section_heading(document, "V. Results")
    add_subheading(document, "A. Detector accuracy")
    add_body(
        document,
        "Table I and Fig. 2 show that fine-tuning YOLO26m on SportsMOT raises mAP50–95 from 0.736 to 0.831 and recall from 0.903 to 0.960 relative to the same pretrained scale. YOLO26n is the fastest detector-only profile but loses 0.241 mAP50–95 relative to the fine-tuned medium model. The result supports task-specific weights for football and smaller generic weights for latency-oriented open deployments.",
    )
    add_figure(
        document,
        repo_root / "docs/benchmarks/research_final/figures/research_detector_quality.png",
        "Fig. 2. Detector quality and throughput under the same SportsMOT validation protocol.",
    )

    add_subheading(document, "B. SportsMOT tracker comparison")
    add_body(
        document,
        "Table II and Fig. 3 show that TrackTrack yields the highest HOTA (71.058) and MOTA (91.511) under shared detections. BoT-SORT ReID has the fewest official identity switches (895) and almost identical IDF1, but its cached throughput is roughly one third of TrackTrack. ByteTrack and FastTracker are substantially faster yet lose more than 12 HOTA points. The result motivates TrackTrack for the quality profile and ByteTrack for the high-throughput profile rather than claiming one tracker dominates every operating point.",
    )
    tracker_rows = []
    for item in research["sportsmot_trackers"]:
        tracker_rows.append(
            [
                item["display_name"],
                f"{item['HOTA']:.1f}",
                f"{item['MOTA']:.1f}",
                f"{item['IDF1']:.1f}",
                str(item["IDSW"]),
                f"{item['cached_pipeline_fps']:.1f}",
            ]
        )
    add_table(
        document,
        "Table II. SportsMOT tracking on 30 validation sequences",
        ["Tracker", "HOTA", "MOTA", "IDF1", "IDSW", "FPS"],
        tracker_rows,
        [1.12, 0.38, 0.40, 0.40, 0.42, 0.36],
    )
    add_figure(
        document,
        repo_root / "docs/benchmarks/research_final/figures/research_tracker_tradeoff.png",
        "Fig. 3. SportsMOT identity-quality and speed trade-off. Values use shared detections.",
    )

    add_subheading(document, "C. Ignore-aware traffic tracking")
    add_body(
        document,
        "Table III and Fig. 4 show that, after correcting ignore-region handling, YOLO26s plus TrackTrack CNN-ReID is the strongest traffic quality profile: 87.605 HOTA, 89.828 AssA, 96.399 IDF1, and zero IDSW. The same detector with ByteTrack reaches 49.52 processed frames per second and remains above 83 HOTA. Increasing the detector to YOLO26m reduces HOTA because this generic checkpoint misses more vehicles on the selected sequence; model size alone is therefore not a quality guarantee.",
    )
    traffic_rows = []
    for item in traffic["tracking"]:
        traffic_rows.append(
            [
                item["short_label"],
                f"{item['HOTA']:.1f}",
                f"{item['IDF1']:.1f}",
                str(item["IDSW"]),
                str(item["FN"]),
                f"{item['processing_fps']:.1f}",
            ]
        )
    add_table(
        document,
        "Table III. Ignore-aware UA-DETRAC detector-tracker matrix",
        ["Profile", "HOTA", "IDF1", "IDSW", "FN", "FPS"],
        traffic_rows,
        [0.82, 0.38, 0.40, 0.42, 0.38, 0.36],
    )
    add_figure(
        document,
        repo_root / "docs/benchmarks/traffic_quality/figures/tracking_identity_speed_tradeoff.png",
        "Fig. 4. Traffic operating points. The quality profile eliminates IDSW; the speed profile exceeds the 30-frame/s source rate.",
    )

    add_subheading(document, "D. Cross-domain transfer")
    add_body(
        document,
        "Table IV shows that the open-vocabulary wildlife profile is detector-limited: zebra tracking obtains 37.719 HOTA and 272 identity switches. In microscopy, the domain adapter yields 67.890 HOTA without preprocessing; contrast-limited adaptive histogram equalization (CLAHE) raises HOTA to 69.526 and IDF1 to 88.009 while reducing foreground throughput from 21.26 to 17.76 processed frames per second. These results reject a universal-preprocessing claim and instead support per-domain detector and preprocessor validation.",
    )
    multidomain_rows = []
    for item in research["multidomain_tracking"]:
        if item["id"].startswith("ua_detrac"):
            continue
        multidomain_rows.append(
            [
                item["domain"],
                str(item["frame_count"]),
                f"{item['HOTA']:.1f}",
                f"{item['MOTA']:.1f}",
                f"{item['IDF1']:.1f}",
                str(item["IDSW"]),
                f"{item['processing_fps']:.1f}",
            ]
        )
    add_table(
        document,
        "Table IV. Cross-domain tracking with official or normalized ground truth",
        ["Domain", "Frames", "HOTA", "MOTA", "IDF1", "IDSW", "FPS"],
        multidomain_rows,
        [0.72, 0.37, 0.36, 0.38, 0.38, 0.38, 0.35],
    )
    semantic = traffic["semantic"]
    add_subheading(document, "E. Semantic quality and cost")
    add_body(
        document,
        "Table V and Fig. 5 report semantic quality and cost. The locate-then-reason worker processed 24 predicted tracks. Only eight can be scored against the class-incomplete UA-DETRAC annotations; the remaining 16 are reported as unscored rather than converted into false semantic errors. Parent-class accuracy is 87.5%, macro-averaged F1 is 85.5%, and coverage is 100%. Fine-label accuracy and unknown-rejection F1 are not measurable because the dataset supplies neither vehicle subtype and color labels nor reviewed unknowns. The complete worker takes 612.28 s, of which 82.09 s is LocateAnything and 530.20 s is Qwen plus overhead.",
    )
    add_table(
        document,
        "Table V. LocateAnything→Qwen semantic evaluation",
        ["Processed", "Scored", "Acc.", "Macro-F1", "Coverage", "Halluc.", "s/track"],
        [
            [
                str(semantic["processed_track_count"]),
                str(semantic["scored_track_count"]),
                f"{100 * semantic['semantic_accuracy']:.1f}%",
                f"{100 * semantic['semantic_macro_f1']:.1f}%",
                f"{100 * semantic['semantic_coverage']:.1f}%",
                f"{100 * semantic['semantic_hallucination_rate']:.1f}%",
                f"{semantic['seconds_per_processed_track']:.2f}",
            ]
        ],
        [0.48, 0.40, 0.40, 0.55, 0.52, 0.52, 0.48],
    )
    add_figure(
        document,
        repo_root / "docs/benchmarks/traffic_quality/figures/semantic_quality_and_cost.png",
        "Fig. 5. Semantic quality, coverage, and per-track cost. Missing fine-label ground truth is shown as not applicable.",
    )

    add_subheading(document, "F. Runtime and webcam behavior")
    add_body(
        document,
        "Table VI summarizes three repeated foreground runs and shows that latency depends on detector scale, object density, rendering, and preprocessing. The street-traffic profile averages 30.05 processed frames per second, while the quality-oriented football profile averages 11.03. Table VII shows that a physical webcam test sustains the 30.0-frame-per-second source rate with 95th-percentile latency between 32.18 and 33.48 ms and at most 0.33% frame drop. These foreground measurements exclude Qwen generation; semantic jobs run asynchronously and update future frames.",
    )
    runtime_rows = []
    selected_runtime = {
        "football_tracktrack_300",
        "traffic_ua_auto_300",
        "traffic_street_auto_300",
        "classroom_auto_300",
        "wildlife_yoloe_300",
        "microscopy_adapter_92",
    }
    for item in research["runtime"]["profiles"]:
        if item["profile"] not in selected_runtime:
            continue
        runtime_rows.append(
            [
                item["profile"].replace("_300", "").replace("_92", ""),
                f"{item['processing_fps_mean']:.1f}",
                f"{item['p95_latency_ms_mean']:.1f}",
                f"{100 * item['drop_rate_mean']:.1f}%",
            ]
        )
    add_table(
        document,
        "Table VI. Repeated foreground runtime by domain",
        ["Profile", "FPS", "p95 ms", "Drop"],
        runtime_rows,
        [1.45, 0.48, 0.48, 0.46],
    )
    webcam_rows = []
    for item in research["physical_webcam"]["profiles"]:
        webcam_rows.append(
            [
                item["profile"]
                .replace("_semantic_deferred", "+semantic")
                .replace("bounded_tracking_only", "tracking"),
                f"{item['source_progress_fps_mean']:.2f}",
                f"{item['p95_latency_ms_mean']:.2f}",
                f"{100 * item['drop_rate_mean']:.2f}%",
            ]
        )
    add_table(
        document,
        "Table VII. Physical webcam foreground measurements",
        ["Profile", "Source FPS", "p95 ms", "Drop"],
        webcam_rows,
        [1.35, 0.58, 0.50, 0.48],
    )

    add_section_heading(document, "VI. Error Analysis and Ablation")
    add_subheading(document, "A. Association and identity failures")
    add_body(
        document,
        "Official IDSW remains the benchmark count. Fig. 6 and Table VIII present a separate heuristic diagnostic that partitions candidate events into fragmentation, identity swap, re-identification failure, association error, and appearance confusion. For TrackTrack on SportsMOT, the diagnostic assigns 37.9% of 1,302 recomputed events to ReID failure and 29.2% to identity swap; for BoT-SORT ReID, the corresponding shares are 45.4% and 33.9%. These labels guide inspection but are not ground-truth causes and must not be ranked as official metrics.",
    )
    add_figure(
        document,
        repo_root / "docs/benchmarks/research_final/figures/research_idsw_taxonomy.png",
        "Fig. 6. Heuristic diagnostic composition of identity failures. Official TrackEval IDSW is reported separately.",
    )
    idsw_rows = []
    for item in research["idsw_taxonomy"]:
        idsw_rows.append(
            [
                item["tracker"],
                str(item["official_idsw"]),
                f"{item['fragmentation_percent']:.1f}",
                f"{item['identity_swap_percent']:.1f}",
                f"{item['re_identification_failure_percent']:.1f}",
                f"{item['association_error_percent']:.1f}",
                f"{item['appearance_confusion_percent']:.1f}",
            ]
        )
    add_table(
        document,
        "Table VIII. Official IDSW and heuristic diagnostic shares (%)",
        ["Tracker", "IDSW", "Frag.", "Swap", "ReID", "Assoc.", "Appear."],
        idsw_rows,
        [0.78, 0.38, 0.38, 0.38, 0.38, 0.40, 0.42],
    )
    add_subheading(document, "B. Class gating")
    add_body(
        document,
        "A controlled UA-DETRAC ablation compares class-agnostic association plus temporal class stabilization against a hard detector-class gate. The class-agnostic policy improves HOTA by 1.134, IDF1 by 1.019, and reduces IDSW from four to three in the pre-correction profile. The selected system therefore keeps semantic class outside the hard geometric association decision unless the task explicitly requires category-isolated tracking.",
    )
    add_subheading(document, "C. Failure modes")
    add_body(
        document,
        "The dominant failures are detector misses under scale change or occlusion, identity fragmentation after long absence, blurred or truncated semantic crops, and vocabulary mismatch. An apparent semantic failure may also be unscorable when the official dataset omits a visible object. For realtime use, fast fly-by objects may leave before rich labels are returned; the system still preserves their base detection and track record, while the deep label may arrive only in the log or replay.",
    )

    add_section_heading(document, "VII. Discussion")
    add_body(
        document,
        "The experiments support two explicit operating points. The quality profile uses YOLO26s and TrackTrack CNN-ReID when identity continuity matters more than 30-frame/s processing. The speed profile uses YOLO26s and ByteTrack when the foreground path must exceed the source rate. In both cases, semantic inference remains asynchronous. Calling the entire stack realtime without this distinction would hide a 25.51-s per-track semantic delay behind a foreground FPS number.",
    )
    add_body(
        document,
        "The study also clarifies what the VLM does not do. It does not replace the detector on every frame, it does not rewrite MOT identities, and it does not guarantee fine-grained accuracy without reviewed fine-label ground truth. Its value is evidence-based semantic enrichment, vocabulary extension, and uncertainty-aware abstention. LocateAnything improves spatial focus before Qwen reasoning, but it does not itself establish temporal identity.",
    )
    add_body(
        document,
        "Table IX places the results beside published SportsMOT values while preserving protocol differences. The literature reports, for example, 64.31 HOTA for ByteTrack, 70.44 for OC-SORT, and 73.91 for BoT-SORT on a custom-detection validation setup [20], whereas our tracker table uses one shared detector cache. These values provide context, not a direct leaderboard claim.",
    )
    add_table(
        document,
        "Table IX. SportsMOT literature context under a different detector protocol",
        ["Published method", "Split", "HOTA", "Detector policy"],
        [
            ["ByteTrack [20]", "val", "64.31", "custom YOLOX"],
            ["OC-SORT [20]", "val", "70.44", "custom YOLOX"],
            ["BoT-SORT [20]", "val", "73.91", "custom YOLOX + ReID"],
            ["Player locating [20]", "val", "86.04", "method-specific"],
        ],
        [1.03, 0.36, 0.42, 1.10],
    )

    add_section_heading(document, "VIII. Limitations")
    add_body(
        document,
        "The corrected traffic matrix covers one 750-frame sequence; broader traffic conclusions require additional official sequences. Fine-grained semantic labels for vehicle subtype, color, classroom role, and wildlife species are not yet supported by reviewed ground truth at sufficient scale. The IDSW taxonomy is heuristic. The open-vocabulary wildlife detector remains weak. Runtime repetitions characterize this RTX 4060 laptop configuration but are not hardware-independent confidence intervals. Finally, 8-bit inference reduces memory but does not make a 4B VLM frame-synchronous.",
    )

    add_section_heading(document, "IX. Conclusion")
    add_body(
        document,
        "A practical semantic tracker on constrained hardware benefits from separating identity continuity from language reasoning. The proposed locate-then-reason architecture preserves an immediate detector-and-tracker output, schedules bounded track evidence, verifies spatial evidence with LocateAnything, and fuses Qwen labels over time with rejection. The strongest measured traffic profile reaches 87.605 HOTA, 96.399 IDF1, and zero IDSW; the speed profile reaches 49.52 frames/s. The results establish a reproducible baseline and, equally importantly, delimit what remains unverified before fine-grained semantic claims can be generalized across domains.",
    )

    add_section_heading(document, "References")
    references = [
        "Ultralytics, “Ultralytics YOLO26: Unified Real-Time End-to-End Vision Models,” arXiv:2606.03748, 2026.",
        "P. Wang et al., “YOLOE: Real-Time Seeing Anything,” arXiv:2503.07465, 2025.",
        "Y. Cui et al., “SportsMOT: A Large Multi-Object Tracking Dataset in Multiple Sports Scenes,” in Proc. ICCV, 2023, pp. 9921–9931.",
        "K. Shim, K. Ko, Y. Yang, and C. Kim, “Focusing on Tracks for Online Multi-Object Tracking,” in Proc. CVPR, 2025, pp. 11687–11696.",
        "A. Bewley, Z. Ge, L. Ott, F. Ramos, and B. Upcroft, “Simple Online and Realtime Tracking,” in Proc. ICIP, 2016, pp. 3464–3468.",
        "N. Wojke, A. Bewley, and D. Paulus, “Simple Online and Realtime Tracking with a Deep Association Metric,” in Proc. ICIP, 2017, pp. 3645–3649.",
        "Y. Zhang et al., “ByteTrack: Multi-Object Tracking by Associating Every Detection Box,” in Proc. ECCV, 2022, pp. 1–21.",
        "N. Aharon, R. Orfaig, and B.-Z. Bobrovsky, “BoT-SORT: Robust Associations Multi-Pedestrian Tracking,” arXiv:2206.14651, 2022.",
        "J. Cao et al., “Observation-Centric SORT: Rethinking SORT for Robust Multi-Object Tracking,” in Proc. CVPR, 2023, pp. 9686–9696.",
        "G. Maggiolino et al., “Deep OC-SORT: Multi-Pedestrian Tracking by Adaptive Re-Identification,” in Proc. ICIP, 2023.",
        "J. Luiten et al., “HOTA: A Higher Order Metric for Evaluating Multi-Object Tracking,” Int. J. Comput. Vis., vol. 129, pp. 548–578, 2021.",
        "J. Luiten and A. Hoffhues, “TrackEval: A Standardized Multi-Object Tracking Evaluation Codebase,” GitHub repository, 2020–2026.",
        "K. Zhou et al., “Omni-Scale Feature Learning for Person Re-Identification,” in Proc. ICCV, 2019, pp. 3702–3712.",
        "S. Bai et al., “Qwen3-VL Technical Report,” arXiv:2511.21631, 2025.",
        "S. Wang et al., “LocateAnything: Fast and High-Quality Vision-Language Grounding with Parallel Box Decoding,” arXiv:2605.27365, 2026.",
        "L. Wen et al., “UA-DETRAC: A New Benchmark and Protocol for Multi-Object Detection and Tracking,” Comput. Vis. Image Underst., vol. 193, 2020.",
        "L. Zhang et al., “AnimalTrack: A Benchmark for Multi-Animal Tracking in the Wild,” Int. J. Comput. Vis., vol. 131, pp. 496–513, 2023.",
        "V. Ulman et al., “An Objective Comparison of Cell-Tracking Algorithms,” Nat. Methods, vol. 14, pp. 1141–1152, 2017.",
        "T. Chamiti et al., “ReferGPT: Towards Zero-Shot Referring Multi-Object Tracking,” in Proc. CVPR Workshops, 2025, pp. 3888–3897.",
        "A. Cioppa et al., “Individual Locating of Soccer Players from a Single Moving View,” Sensors, vol. 23, no. 18, 2023.",
    ]
    for index, reference in enumerate(references, start=1):
        add_reference(document, index, reference)

    # A trailing continuous section lets Word balance the preceding two-column
    # reference section without adding a visible page break or page number.
    balancing_section = document.add_section(WD_SECTION.CONTINUOUS)
    _configure_section(balancing_section)
    _set_columns(balancing_section, 1)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = TITLE
    document.core_properties.subject = (
        "Semantic multi-object tracking with LocateAnything and Qwen3-VL"
    )
    document.core_properties.author = "Hieu Le Nho"
    document.core_properties.keywords = "MOT, VLM, Qwen3-VL, LocateAnything, YOLO26, TrackTrack"
    document.save(output_docx)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Build the IEEE-style semantic tracking paper DOCX."
    )
    parser.add_argument("--repo-root", type=Path, default=Path(__file__).resolve().parents[2])
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("outputs/reports/research_paper/locate_then_reason_ieee_paper.docx"),
    )
    args = parser.parse_args()
    repo_root = args.repo_root.resolve()
    output = args.output if args.output.is_absolute() else repo_root / args.output
    build_document(repo_root, output.resolve())
    print(json.dumps({"status": "ok", "output": str(output.resolve())}, indent=2))


if __name__ == "__main__":
    main()
